"""多格式匯出 API（TASK-010）。

Word (.docx) 與 PDF 由後端產生；純文字匯出由前端直接產生（不呼叫此 API）。
"""
from __future__ import annotations

import io

from docx import Document
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas

import jobstore
from auth import CurrentUser, get_current_user

router = APIRouter()

_CJK_FONT = "MSung-Light"
pdfmetrics.registerFont(UnicodeCIDFont(_CJK_FONT))


def _load_minutes(job_id: str) -> dict:
    job = jobstore.get_job(job_id)
    if job is None or job.get("minutes") is None:
        raise HTTPException(status_code=404, detail="找不到會議紀錄，請先完成 /summarize")
    return job["minutes"]


def _meeting_info_lines(minutes: dict) -> list[str]:
    """把 meeting_info 轉成顯示用的一行行文字，未提及欄位顯示「未提及」。"""
    info = minutes.get("meeting_info") or {}
    date = info.get("date") or "未提及"
    time = info.get("time") or "未提及"
    location = info.get("location") or "未提及"
    participants = info.get("participants") or []
    participants_text = "、".join(participants) if participants else "未提及"
    return [
        f"日期：{date}",
        f"時間：{time}",
        f"地點：{location}",
        f"參與者：{participants_text}",
    ]


def _build_docx(minutes: dict, job_id: str) -> bytes:
    doc = Document()
    doc.add_heading(f"會議紀錄 - {job_id}", level=1)

    doc.add_heading("會議資訊", level=2)
    for line in _meeting_info_lines(minutes):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("摘要", level=2)
    doc.add_paragraph(minutes.get("summary", ""))

    doc.add_heading("決定事項", level=2)
    decisions = minutes.get("decisions", [])
    if decisions:
        for d in decisions:
            doc.add_paragraph(d, style="List Bullet")
    else:
        doc.add_paragraph("（無）")

    doc.add_heading("待辦事項", level=2)
    action_items = minutes.get("action_items", [])
    if action_items:
        for item in action_items:
            owner = item.get("owner", "-")
            task = item.get("task", "-")
            due = item.get("due", "-")
            doc.add_paragraph(f"[{owner}] {task}（期限：{due}）", style="List Bullet")
    else:
        doc.add_paragraph("（無）")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _build_pdf(minutes: dict, job_id: str) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    y = height - 60

    def _line(text: str, size: int = 12, gap: int = 20):
        nonlocal y
        c.setFont(_CJK_FONT, size)
        c.drawString(50, y, text)
        y -= gap
        if y < 60:
            c.showPage()
            y = height - 60

    _line(f"會議紀錄 - {job_id}", size=16, gap=30)

    _line("會議資訊", size=14, gap=22)
    for line in _meeting_info_lines(minutes):
        _line(line)

    _line("摘要", size=14, gap=22)
    for chunk in _wrap(minutes.get("summary", ""), 40):
        _line(chunk)

    _line("決定事項", size=14, gap=22)
    decisions = minutes.get("decisions", [])
    if decisions:
        for d in decisions:
            _line(f"- {d}")
    else:
        _line("（無）")

    _line("待辦事項", size=14, gap=22)
    action_items = minutes.get("action_items", [])
    if action_items:
        for item in action_items:
            owner = item.get("owner", "-")
            task = item.get("task", "-")
            due = item.get("due", "-")
            _line(f"- [{owner}] {task}（期限：{due}）")
    else:
        _line("（無）")

    c.save()
    buf.seek(0)
    return buf.read()


def _wrap(text: str, width: int) -> list[str]:
    if not text:
        return [""]
    return [text[i : i + width] for i in range(0, len(text), width)]


@router.get("/export/{job_id}")
async def export_meeting(
    job_id: str, format: str = "docx", user: CurrentUser = Depends(get_current_user)
):
    """匯出指定 job 的會議紀錄。format 支援 docx / pdf。"""
    minutes = _load_minutes(job_id)

    if format == "docx":
        content = _build_docx(minutes, job_id)
        media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        filename = f"{job_id}.docx"
    elif format == "pdf":
        content = _build_pdf(minutes, job_id)
        media_type = "application/pdf"
        filename = f"{job_id}.pdf"
    else:
        raise HTTPException(status_code=400, detail="format 僅支援 docx 或 pdf")

    return StreamingResponse(
        io.BytesIO(content),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
